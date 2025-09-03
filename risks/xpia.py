"""
Run from the risks directory:
    uv run xpia.py -t streamable-http
"""

from mcp.server.fastmcp import FastMCP
import os
import argparse

mcp = FastMCP("xpia")

SENSITIVE_KEYWORDS = {"confidential", "secret", "passwords", "keys"}
SENSITIVE_DENYLIST = set()


def _init_sensitive_denylist(base_dir: str) -> None:
  """Initialize explicit sensitive file denylist."""
  # Explicitly deny known sensitive files
  defaults = [
      os.path.join(base_dir, "data", "confidential.txt"),
  ]
  for p in defaults:
    SENSITIVE_DENYLIST.add(os.path.realpath(p))

  # Optional: add more via env var (colon-separated relative paths under base_dir)
  extra = os.environ.get("XPIA_SENSITIVE_DENYLIST", "")
  if extra:
    for rel in [x.strip() for x in extra.split(os.pathsep) if x.strip()]:
      SENSITIVE_DENYLIST.add(os.path.realpath(os.path.join(base_dir, rel)))


def _is_sensitive_path(abs_path: str, base_dir: str) -> tuple[bool, str]:
  """Return (blocked, reason) if the path should be blocked."""
  rp = os.path.realpath(abs_path)

  # Must remain within base_dir even across symlinks
  try:
    if os.path.commonpath([rp, os.path.realpath(base_dir)]) != os.path.realpath(base_dir):
      return True, "Path escapes application directory."
  except Exception:
    # In case commonpath errors on different drives, fail closed
    return True, "Unable to validate path location."

  # Explicit denylist
  if rp in SENSITIVE_DENYLIST:
    return True, "This file is on the sensitive denylist."

  # Simple filename keyword heuristic
  fname = os.path.basename(rp).lower()
  if any(k in fname for k in SENSITIVE_KEYWORDS):
    return True, "Filename matches sensitive keywords."

  return False, ""


# Initialize denylist once
_INIT_BASE_DIR = os.path.abspath(os.path.dirname(__file__))
_init_sensitive_denylist(_INIT_BASE_DIR)

def _extract_text_from_pdf(abs_filepath: str) -> str:
  """Extract text from a PDF using pypdf with simple, clear handling."""
  try:
    from pypdf import PdfReader
  except Exception:
    return "Error: PDF support requires the 'pypdf' package. Install with: uv add pypdf"

  try:
    reader = PdfReader(abs_filepath)
    if getattr(reader, "is_encrypted", False):
      return "Error: Encrypted PDF is not supported."

    parts = []
    for i, page in enumerate(reader.pages):
      try:
        text = (page.extract_text() or "").strip()
      except Exception:
        text = ""
      if text:
        parts.append(f"--- Page {i+1} ---\n{text}")

    return "\n\n".join(parts) if parts else "Warning: No extractable text found in PDF."
  except FileNotFoundError:
    return f"Error: File not found: {abs_filepath}"
  except PermissionError:
    return f"Error: Permission denied for file: {abs_filepath}"
  except Exception as e:
    return f"Error reading PDF: {e}"


def _extract_text_from_docx(abs_filepath: str) -> str:
  """Extract text from a DOCX using python-docx with simple handling."""
  try:
    from docx import Document
  except Exception:
    return "Error: DOCX support requires the 'python-docx' package. Install with: uv add python-docx"

  try:
    doc = Document(abs_filepath)
    parts = []

    # Paragraphs
    for p in doc.paragraphs:
      t = p.text.strip()
      if t:
        parts.append(t)

    # Tables (best-effort)
    try:
      for table in doc.tables:
        for row in table.rows:
          row_text = [cell.text.strip()
                      for cell in row.cells if cell.text.strip()]
          if row_text:
            parts.append(" | ".join(row_text))
    except Exception:
      pass

    return "\n".join(parts) if parts else "Warning: No extractable text found in DOCX."
  except FileNotFoundError:
    return f"Error: File not found: {abs_filepath}"
  except PermissionError:
    return f"Error: Permission denied for file: {abs_filepath}"
  except Exception as e:
    return f"Error reading DOCX: {e}"


@mcp.tool()
def read_file(filepath: str) -> str:
  """Read the content of a local file (.txt, .pdf, or .docx) and return it as a string."""
  try:
    abs_filepath = os.path.realpath(os.path.abspath(filepath))
    base_dir = _INIT_BASE_DIR

    # Allow only files within the application directory or subdirectories
    if os.path.commonpath([abs_filepath, base_dir]) != base_dir:
      return "Error: Cannot access files outside the application directory for security reasons."

    # Block sensitive paths/filenames
    blocked, reason = _is_sensitive_path(abs_filepath, base_dir)
    if blocked:
      return f"Blocked by server policy: {reason}"

    ext = os.path.splitext(abs_filepath)[1].lower()
    if ext == ".pdf":
      return _extract_text_from_pdf(abs_filepath)
    elif ext == ".docx":
      return _extract_text_from_docx(abs_filepath)
    elif ext == ".txt":
      with open(abs_filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()
    else:
      return "Error: Unsupported file type. Only .txt, .pdf, and .docx are supported."

  except FileNotFoundError:
    return f"Error: File not found: {filepath}"
  except PermissionError:
    return f"Error: Permission denied for file: {filepath}"
  except Exception as e:
    return f"Error reading file: {e}"


if __name__ == "__main__":
  parser = argparse.ArgumentParser(description="XPIA")
  parser.add_argument(
      "--transport", "-t", choices=["stdio", "sse", "streamable-http"], default="stdio")
  args = parser.parse_args()
  mcp.run(transport=args.transport)